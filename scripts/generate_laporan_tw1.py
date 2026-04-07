"""
Generate Laporan Progress TW1 2026 - Laboratorium Inovasi Digital
Kepala Lab: Aidil Saputra Kirsan
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "..", "Laporan_Progress_TW1_2026_LabInovasiDigital_v3.docx")


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        tag = OxmlElement(f"w:{edge}")
        for attr, val in attrs.items():
            tag.set(qn(f"w:{attr}"), val)
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p


def add_paragraph(doc, text, bold=False, italic=False, size=11, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    return p


def add_placeholder_box(doc, label):
    """Menambahkan kotak placeholder untuk screenshot."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(14)
    set_cell_bg(cell, "E8F4F8")
    set_cell_border(cell,
        top={"val": "single", "sz": "6", "color": "2E86AB"},
        bottom={"val": "single", "sz": "6", "color": "2E86AB"},
        left={"val": "single", "sz": "6", "color": "2E86AB"},
        right={"val": "single", "sz": "6", "color": "2E86AB"},
    )
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"📷  {label}")
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    run.font.size = Pt(10)
    run.font.italic = True

    # spacer bawah
    p2 = cell.add_paragraph("[ Sisipkan screenshot di sini ]")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p2.runs[0].font.size = Pt(9)

    doc.add_paragraph()  # spasi setelah tabel


def add_info_table(doc, rows_data):
    """Tabel header info laporan."""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(rows_data):
        r = table.rows[i]
        c0, c1 = r.cells[0], r.cells[1]
        c0.width = Cm(5)
        c1.width = Cm(11)
        set_cell_bg(c0, "1B4F72")
        p0 = c0.paragraphs[0]
        run0 = p0.add_run(key)
        run0.bold = True
        run0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run0.font.size = Pt(10)

        p1 = c1.paragraphs[0]
        p1.add_run(val).font.size = Pt(10)


def add_section_table(doc, headers, rows, col_widths=None):
    """Tabel data umum."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "1B4F72")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        if col_widths:
            cell.width = Cm(col_widths[i])

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = "EBF5FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if col_widths:
                cell.width = Cm(col_widths[ci])

    doc.add_paragraph()


# ─────────────────────────────────────────────
# MAIN DOCUMENT
# ─────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# Default font
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# ── KOVER / HEADER ──────────────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("LAPORAN PROGRESS")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p_sub.add_run("Indikator Kinerja Utama Kepala Laboratorium Inovasi Digital")
run2.bold = True
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

p_tw = doc.add_paragraph()
p_tw.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p_tw.add_run("Triwulan I (Januari – Maret 2026)")
run3.bold = True
run3.font.size = Pt(13)

doc.add_paragraph()

add_info_table(doc, [
    ("Unit",              "Laboratorium Inovasi Digital"),
    ("Fakultas/Jurusan",  "Jurusan Teknik Elektro, Informatika, dan Bisnis"),
    ("Institusi",         "Institut Teknologi Kalimantan"),
    ("Kepala Lab",        "Aidil Saputra Kirsan"),
    ("Periode Laporan",   "Triwulan I – Tahun 2026 (Januari s.d. Maret 2026)"),
    ("Tanggal Penyusunan", datetime.date.today().strftime("%d %B %Y")),
])

doc.add_page_break()

# ── I. PENDAHULUAN ───────────────────────────────────────────────────────────
add_heading(doc, "I. PENDAHULUAN", 1, (0x1B, 0x4F, 0x72))
add_paragraph(doc,
    "Laboratorium Inovasi Digital merupakan unit laboratorium yang bertugas mendukung "
    "kegiatan Penelitian dan Pengabdian kepada Masyarakat (PPM) dosen di lingkungan "
    "Jurusan Teknik Elektro, Informatika, dan Bisnis, Institut Teknologi Kalimantan. "
    "Pada periode ini, laboratorium secara aktif melayani dua program studi, yaitu Sistem "
    "Informasi (SI) dan Bisnis Digital (BD), dengan total 25 dosen yang datanya telah "
    "tercatat dan terpantau secara berkala."
)
doc.add_paragraph()
add_paragraph(doc,
    "Laporan ini menyajikan perkembangan pelaksanaan kegiatan pengembangan laboratorium "
    "pada Triwulan I Tahun 2026 (Januari s.d. Maret 2026) sesuai dengan Indikator Kinerja "
    "Utama (IKU) yang telah ditetapkan, meliputi: (1) Klaster dan Roadmap PPM, "
    "(2) Sistem Penunjang PPM, (3) Galeri/Katalog Inovasi, (4) Workshop Pengembangan Skill, "
    "dan (5) Seminar Ilmiah."
)

doc.add_page_break()

# ── II. CAKUPAN DATA & SISTEM ────────────────────────────────────────────────
add_heading(doc, "II. CAKUPAN DATA DAN SISTEM", 1, (0x1B, 0x4F, 0x72))
add_paragraph(doc,
    "Seluruh data kegiatan PPM yang ditampilkan dan dikelola oleh Laboratorium Inovasi "
    "Digital bersumber dari profil resmi dosen yang terdaftar pada SINTA (Sistem Indeksasi "
    "Penelitian Nasional milik Kemdiktisaintek). Berikut adalah cakupan data yang telah "
    "tercatat dalam sistem laboratorium:"
)
doc.add_paragraph()

add_section_table(doc,
    ["Program Studi", "Jurusan", "Jumlah Dosen", "Status Data"],
    [
        ["Sistem Informasi (SI)", "Teknik Elektro, Informatika, dan Bisnis", "15 dosen", "Aktif & Terverifikasi"],
        ["Bisnis Digital (BD)",   "Teknik Elektro, Informatika, dan Bisnis", "10 dosen", "Aktif & Terverifikasi"],
    ],
    col_widths=[5, 7, 3.5, 4.5]
)

add_paragraph(doc,
    "Data diperbaharui secara berkala dengan mengambil informasi terbaru langsung dari "
    "laman resmi SINTA setiap dosen. Pembaruan data terakhir dilakukan pada 6 Februari 2026."
)

doc.add_page_break()

# ── III. PROGRESS PER INDIKATOR ─────────────────────────────────────────────
add_heading(doc, "III. PROGRESS PELAKSANAAN KEGIATAN", 1, (0x1B, 0x4F, 0x72))

# ─── 3.1 Klaster dan Roadmap PPM ────────────────────────────────────────────
add_heading(doc, "3.1  Klaster dan Roadmap PPM", 2, (0x21, 0x81, 0x8B))

add_paragraph(doc, "A. Deskripsi Kegiatan", bold=True)
add_paragraph(doc,
    "Kegiatan ini bertujuan untuk memetakan dan mengelompokkan topik-topik penelitian dan "
    "pengabdian yang selama ini telah dilakukan oleh para dosen, sehingga tersedia gambaran "
    "yang jelas mengenai kekuatan dan arah riset laboratorium. Peta pengelompokan ini "
    "diharapkan dapat menjadi dasar perencanaan kolaborasi antar dosen maupun antar program studi."
)

doc.add_paragraph()
add_paragraph(doc, "B. Progress yang Dicapai", bold=True)

add_paragraph(doc, "1) Pengelompokan (Klaster) Topik PPM", bold=True, indent=True)
add_bullet(doc, "Sebanyak 148 judul kegiatan Penelitian dan Pengabdian dosen Prodi Sistem Informasi berhasil dianalisis dan dikelompokkan")
add_bullet(doc, "Terbentuk 12 kelompok topik (klaster) yang mencerminkan bidang-bidang unggulan dosen")
add_bullet(doc, "Setiap kelompok menampilkan daftar dosen yang terlibat, sehingga memudahkan identifikasi peluang kolaborasi")

doc.add_paragraph()
add_paragraph(doc, "Tabel 1. Lima Kelompok Topik PPM Terbesar – Prodi Sistem Informasi", italic=True)
add_section_table(doc,
    ["No", "Kelompok Topik", "Tema Utama", "Jumlah Karya"],
    [
        ["1", "Pemanfaatan Teknologi untuk Masyarakat", "Pemanfaatan teknologi digital, peningkatan kapasitas", "27"],
        ["2", "Pemberdayaan Wilayah Balikpapan",        "Pengembangan kelurahan, kota, dan optimalisasi layanan","25"],
        ["3", "Pengembangan Sistem Layanan Digital",    "Sistem informasi layanan publik dan administrasi",      "17"],
        ["4", "Sistem Cerdas dan Rekayasa Perangkat Lunak", "Sistem berbasis kecerdasan, rekayasa lingkungan",  "16"],
        ["5", "Teknologi Pembelajaran dan Media",       "E-learning, media pembelajaran, adopsi teknologi",     "16"],
    ],
    col_widths=[1, 5.5, 6, 3.5]
)

add_paragraph(doc, "2) Peta Arah Riset (Roadmap PPM)", bold=True, indent=True)
add_bullet(doc, "Peta arah riset mencakup data selama 7 tahun, dari tahun 2018 hingga 2025")
add_bullet(doc, "Ditampilkan dalam bentuk diagram alur yang menggambarkan bagaimana topik-topik riset berkembang dari tahun ke tahun")
add_bullet(doc, "Topik yang paling banyak dikerjakan dosen pada tahun 2025: Digitalisasi (20 karya), Kawasan IKN (10 karya), dan Pemanfaatan Teknologi (9 karya)")
add_bullet(doc, "Jumlah kegiatan PPM tumbuh pesat: dari 1 karya pada 2018 menjadi 78 karya pada 2025")

doc.add_paragraph()
add_paragraph(doc, "Tabel 2. Perkembangan Jumlah Kegiatan PPM per Tahun", italic=True)
add_section_table(doc,
    ["Tahun", "2018", "2019", "2020", "2021", "2022", "2023", "2024", "2025"],
    [["Jumlah Kegiatan", "1", "10", "11", "12", "17", "44", "55", "78"]],
    col_widths=[3.5, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8, 1.8]
)

doc.add_paragraph()
add_paragraph(doc, "C. Bukti Pelaksanaan", bold=True)
add_placeholder_box(doc, "Tampilan Halaman Pengelompokan Topik PPM – 12 Kelompok Riset Dosen")
add_placeholder_box(doc, "Tampilan Peta Arah Riset (Roadmap) – Alur Topik PPM 2018–2025")

add_paragraph(doc, "D. Kendala", bold=True)
add_bullet(doc, "Pengelompokan topik baru mencakup data Prodi Sistem Informasi; data Prodi Bisnis Digital masih dalam proses penyesuaian agar dapat digabungkan")
add_bullet(doc, "Peluang kolaborasi antara kedua prodi belum dapat ditampilkan karena proses penggabungan data masih berlangsung")

doc.add_paragraph()
add_paragraph(doc, "E. Strategi Tindak Lanjut", bold=True)
add_bullet(doc, "Menggabungkan data Prodi Bisnis Digital agar peta kolaborasi lintas prodi dapat segera tersedia")
add_bullet(doc, "Secara bertahap memperluas cakupan ke prodi lain di jurusan (Teknik Elektro, Informatika, Teknik Biomedis) pada TW2–TW3 2026")
add_bullet(doc, "Mengadakan sesi diskusi bersama dosen untuk memastikan pengelompokan topik sesuai dengan bidang keahlian masing-masing")

doc.add_page_break()

# ─── 3.2 Sistem Penunjang PPM ────────────────────────────────────────────────
add_heading(doc, "3.2  Sistem Penunjang PPM", 2, (0x21, 0x81, 0x8B))

add_paragraph(doc, "A. Deskripsi Kegiatan", bold=True)
add_paragraph(doc,
    "Laboratorium Inovasi Digital telah membangun dan menjalankan sebuah sistem informasi "
    "berbasis web yang dapat diakses secara terbuka oleh seluruh civitas akademika. Sistem "
    "ini dirancang untuk memudahkan pemantauan dan pengelolaan data kegiatan PPM dosen "
    "secara terpusat dan real-time."
)

doc.add_paragraph()
add_paragraph(doc, "B. Progress yang Dicapai", bold=True)
add_paragraph(doc, "Fitur-fitur yang telah aktif berjalan dalam sistem:", italic=True)

add_section_table(doc,
    ["No", "Fitur", "Fungsi", "Status"],
    [
        ["1", "Pantau Kinerja Dosen",      "Menampilkan rekap publikasi, jumlah sitasi, dan produktivitas Tridarma seluruh dosen dari 2 prodi (25 dosen)",  "✅ Aktif"],
        ["2", "Pantau Dana Hibah",         "Memantau perolehan dan tren dana hibah penelitian maupun pengabdian setiap dosen",                              "✅ Aktif"],
        ["3", "Pencarian Pakar",           "Memudahkan pencarian dosen yang sesuai berdasarkan topik atau bidang keahlian tertentu",                        "✅ Aktif"],
        ["4", "Data Akreditasi (DTPS)",    "Menyediakan data dan perhitungan rasio dosen untuk keperluan penyusunan borang akreditasi",                     "✅ Aktif"],
        ["5", "Dukungan Multi Prodi",      "Sistem dirancang agar dapat diperluas untuk menampung data dari prodi-prodi lain di jurusan",                   "✅ Aktif (2 prodi)"],
    ],
    col_widths=[1, 4, 8, 3]
)

doc.add_paragraph()
add_paragraph(doc, "C. Bukti Pelaksanaan", bold=True)
add_placeholder_box(doc, "Tampilan Halaman Pantau Kinerja Dosen – Rekap Statistik Prodi SI dan BD")
add_placeholder_box(doc, "Tampilan Fitur Pencarian Pakar – Contoh Hasil Pencarian Bidang Keahlian")
add_placeholder_box(doc, "Tampilan Halaman Pantau Dana Hibah – Grafik Tren Perolehan Hibah Dosen")

add_paragraph(doc, "D. Kendala", bold=True)
add_bullet(doc, "Pembaruan data saat ini masih dilakukan secara manual dan belum berjalan otomatis sesuai jadwal")
add_bullet(doc, "Layanan pengambilan data dari Google Scholar memiliki batasan kuota sehingga perlu dikelola dengan hati-hati")
add_bullet(doc, "Cakupan data masih terbatas pada 2 prodi; prodi-prodi lain di jurusan belum terintegrasi")

doc.add_paragraph()
add_paragraph(doc, "E. Strategi Tindak Lanjut", bold=True)
add_bullet(doc, "Menjadwalkan pembaruan data secara otomatis agar informasi di sistem selalu mutakhir tanpa perlu campur tangan manual")
add_bullet(doc, "Berkoordinasi dengan perwakilan dosen dari prodi-prodi lain untuk proses pendataan secara bertahap")
add_bullet(doc, "Mencari alternatif sumber data tambahan untuk mengurangi ketergantungan pada layanan pihak ketiga yang terbatas")

doc.add_page_break()

# ─── 3.3 Galeri / Katalog Inovasi ────────────────────────────────────────────
add_heading(doc, "3.3  Galeri/Katalog Inovasi", 2, (0x21, 0x81, 0x8B))

add_paragraph(doc, "A. Deskripsi Kegiatan", bold=True)
add_paragraph(doc,
    "Galeri/Katalog Inovasi merupakan ruang pamer digital yang menampilkan seluruh karya "
    "dan hasil kerja dosen Laboratorium Inovasi Digital. Katalog ini dapat diakses oleh "
    "sivitas akademika maupun masyarakat umum, dan berfungsi sebagai wajah laboratorium "
    "dalam mendiseminasikan capaian riset dan pengabdian."
)

doc.add_paragraph()
add_paragraph(doc, "B. Progress yang Dicapai", bold=True)

add_section_table(doc,
    ["Jenis Karya", "Keterangan", "Tersedia"],
    [
        ["Penelitian",       "Karya penelitian dosen yang telah terdaftar dan terindeks secara nasional",     "✅"],
        ["Pengabdian",       "Dokumentasi kegiatan pengabdian kepada masyarakat yang telah dilaksanakan",     "✅"],
        ["Jurnal Internasional (Scopus)", "Artikel dosen yang terbit di jurnal bereputasi internasional",    "✅"],
        ["Jurnal Nasional (SINTA)",       "Artikel dosen yang terbit di jurnal nasional terakreditasi",      "✅"],
        ["Buku",             "Buku ajar maupun buku referensi yang telah diterbitkan",                       "✅"],
        ["Kekayaan Intelektual / Paten",  "Hak cipta, paten, dan kekayaan intelektual lainnya",             "✅"],
    ],
    col_widths=[4.5, 9, 2.5]
)

add_bullet(doc, "Tersedia fitur pencarian dan penyaringan karya berdasarkan tahun, program studi, dan jenis karya")
add_bullet(doc, "Tersedia tampilan layar besar (TV Display Mode) khusus untuk keperluan pameran atau presentasi di ruang laboratorium")
add_bullet(doc, "Data katalog dapat diekspor ke dalam format spreadsheet untuk keperluan pelaporan")

doc.add_paragraph()
add_paragraph(doc, "C. Bukti Pelaksanaan", bold=True)
add_placeholder_box(doc, "Tampilan Galeri/Katalog Inovasi – Daftar Karya Dosen dengan Fitur Filter")
add_placeholder_box(doc, "Tampilan TV Display Mode – Tampilan Pameran Karya di Layar Besar")

add_paragraph(doc, "D. Kendala", bold=True)
add_bullet(doc, "Karya yang ditampilkan di galeri saat ini hanya yang sudah terdaftar di SINTA; karya yang belum didaftarkan dosen ke SINTA belum dapat tampil")
add_bullet(doc, "Kelengkapan informasi tiap karya (seperti ringkasan dan tautan artikel) bergantung pada seberapa lengkap dosen mengisi profil SINTA mereka")

doc.add_paragraph()
add_paragraph(doc, "E. Strategi Tindak Lanjut", bold=True)
add_bullet(doc, "Mengimbau para dosen untuk secara rutin memperbarui dan melengkapi profil SINTA mereka agar informasi di galeri semakin lengkap")
add_bullet(doc, "Menambahkan fitur agar dosen dapat secara mandiri memasukkan karya yang belum terdaftar di SINTA")

doc.add_page_break()

# ─── 3.4 Workshop Pengembangan Skill ─────────────────────────────────────────
add_heading(doc, "3.4  Workshop Pengembangan Skill", 2, (0x21, 0x81, 0x8B))

add_paragraph(doc, "A. Deskripsi Kegiatan", bold=True)
add_paragraph(doc,
    "Kegiatan workshop pengembangan skill bertujuan untuk meningkatkan kemampuan dosen "
    "dan mahasiswa yang berafiliasi dengan Laboratorium Inovasi Digital, baik dalam hal "
    "peningkatan kualitas riset, penguasaan teknologi terkini, maupun pengembangan "
    "kewirausahaan berbasis inovasi digital."
)

doc.add_paragraph()
add_paragraph(doc, "B. Progress yang Dicapai", bold=True)
add_paragraph(doc,
    "Pada Triwulan I 2026, kegiatan workshop pengembangan skill yang diselenggarakan secara "
    "resmi oleh Laboratorium Inovasi Digital belum terlaksana. Kegiatan ini masih dalam "
    "tahap perencanaan dan akan dijadwalkan pada triwulan berikutnya."
)

doc.add_paragraph()
add_paragraph(doc, "C. Bukti Pelaksanaan", bold=True)
add_paragraph(doc, "– Belum ada bukti pelaksanaan pada TW1 2026.", italic=True)

add_paragraph(doc, "D. Kendala", bold=True)
add_bullet(doc, "Belum tersedia jadwal workshop resmi yang diorganisir langsung oleh laboratorium pada TW1 2026")
add_bullet(doc, "Kegiatan pelatihan yang berjalan masih merupakan inisiatif masing-masing dosen, belum terkoordinasi sebagai program bersama laboratorium")
add_bullet(doc, "Padatnya agenda dosen di awal semester menjadi kendala untuk menjadwalkan kegiatan tambahan")

doc.add_paragraph()
add_paragraph(doc, "E. Strategi Tindak Lanjut", bold=True)
add_bullet(doc, "Menyusun kalender kegiatan workshop Laboratorium Inovasi Digital untuk periode TW2–TW4 2026")
add_bullet(doc, "Merencanakan minimal satu kegiatan workshop per triwulan, misalnya pelatihan proposal hibah, penggunaan alat bantu riset, atau literasi data")
add_bullet(doc, "Berkoordinasi dengan Ketua Program Studi SI dan BD agar workshop laboratorium dapat disinergikan dengan kegiatan kemahasiswaan atau akademik")

doc.add_page_break()

# ─── 3.5 Seminar Ilmiah ──────────────────────────────────────────────────────
add_heading(doc, "3.5  Seminar Ilmiah", 2, (0x21, 0x81, 0x8B))

add_paragraph(doc, "A. Deskripsi Kegiatan", bold=True)
add_paragraph(doc,
    "Kegiatan seminar ilmiah bertujuan untuk menyebarluaskan hasil-hasil penelitian dan "
    "pengabdian yang telah dilakukan oleh dosen di lingkungan laboratorium, sekaligus "
    "menjadi ruang diskusi dan pertukaran ide ilmiah antar dosen, mahasiswa, maupun pihak "
    "luar yang berminat."
)

doc.add_paragraph()
add_paragraph(doc, "B. Progress yang Dicapai", bold=True)
add_paragraph(doc,
    "Pada Triwulan I 2026, kegiatan seminar ilmiah yang diselenggarakan secara resmi oleh "
    "Laboratorium Inovasi Digital belum terlaksana. Kegiatan ini masih dalam tahap "
    "perencanaan dan akan dijadwalkan pada triwulan berikutnya."
)

doc.add_paragraph()
add_paragraph(doc, "C. Bukti Pelaksanaan", bold=True)
add_paragraph(doc, "– Belum ada bukti pelaksanaan pada TW1 2026.", italic=True)

add_paragraph(doc, "D. Kendala", bold=True)
add_bullet(doc, "Kegiatan seminar yang diselenggarakan oleh Lab Inovasi Digital belum terlaksana pada TW1 2026")
add_bullet(doc, "Waktu persiapan yang tersedia di awal tahun masih terbatas")
add_bullet(doc, "Penyesuaian jadwal dengan dosen yang akan menjadi narasumber membutuhkan waktu lebih awal")

doc.add_paragraph()
add_paragraph(doc, "E. Strategi Tindak Lanjut", bold=True)
add_bullet(doc, "Menjadwalkan seminar ilmiah perdana Laboratorium Inovasi Digital pada TW2 2026 (April–Juni 2026)")
add_bullet(doc, "Mengundang dosen dari dalam laboratorium yang memiliki karya terbaru untuk tampil sebagai narasumber")
add_bullet(doc, "Membuka peluang kerja sama dengan laboratorium lain di jurusan untuk menyelenggarakan seminar bersama sehingga peserta lebih beragam")
add_bullet(doc, "Mendokumentasikan seluruh kegiatan seminar ke dalam galeri/katalog laboratorium sebagai rekam jejak")

doc.add_page_break()

# ── IV. REKAPITULASI PROGRESS ────────────────────────────────────────────────
add_heading(doc, "IV. REKAPITULASI PROGRESS TW1 2026", 1, (0x1B, 0x4F, 0x72))

add_section_table(doc,
    ["No", "Indikator Kinerja Utama", "Target TW1", "Realisasi", "Status", "Capaian"],
    [
        ["1", "Klaster dan Roadmap PPM",    "Tersedianya peta pengelompokan dan arah riset PPM",  "12 kelompok topik terbentuk; peta arah riset 2018–2025 aktif dan dapat diakses",      "✅ Tercapai",      "100%"],
        ["2", "Sistem Penunjang PPM",        "Sistem informasi penunjang PPM aktif berjalan",      "5 fitur utama aktif; data 25 dosen dari 2 prodi sudah terintegrasi",                  "✅ Tercapai",      "100%"],
        ["3", "Galeri/Katalog Inovasi",      "Galeri digital karya dosen tersedia dan dapat diakses", "6 jenis karya ditampilkan; fitur pencarian dan tampilan pameran sudah aktif",     "✅ Tercapai",      "100%"],
        ["4", "Workshop Pengembangan Skill", "Min. 1 kegiatan workshop terlaksana",                "Belum terlaksana; dalam tahap perencanaan untuk TW2 2026",  "⏳ Direncanakan", "0%"],
        ["5", "Seminar Ilmiah",              "Min. 1 seminar ilmiah terlaksana",                   "Belum terlaksana; dalam tahap perencanaan untuk TW2 2026",  "⏳ Direncanakan", "0%"],
    ],
    col_widths=[1, 4, 4, 5, 2.5, 1.5]
)

doc.add_paragraph()
add_paragraph(doc,
    "Keterangan: Indikator 1 hingga 3 telah tercapai sepenuhnya karena merupakan hasil "
    "pembangunan sistem yang telah dikerjakan dan kini beroperasi penuh sejak TW1 2026. "
    "Indikator 4 dan 5 memerlukan perencanaan dan pelaksanaan kegiatan tatap muka yang "
    "akan diupayakan pada TW2 2026.",
    italic=True
)

doc.add_page_break()

# ── V. PENUTUP ───────────────────────────────────────────────────────────────
add_heading(doc, "V. PENUTUP", 1, (0x1B, 0x4F, 0x72))
add_paragraph(doc,
    "Laboratorium Inovasi Digital telah menunjukkan perkembangan yang nyata pada Triwulan I "
    "Tahun 2026, terutama dalam hal penyediaan sistem informasi penunjang PPM yang kini "
    "sudah dapat digunakan secara aktif. Sistem yang telah tersedia meliputi pemantauan "
    "kinerja dosen, pemetaan topik riset, galeri karya inovasi, serta berbagai fitur pendukung "
    "perencanaan dan pelaporan akademik."
)
doc.add_paragraph()
add_paragraph(doc,
    "Pada triwulan berikutnya, fokus kegiatan akan diarahkan pada: (1) penyelenggaraan "
    "workshop pengembangan skill minimal satu kali, (2) pelaksanaan seminar ilmiah perdana "
    "laboratorium, (3) perluasan cakupan data ke program studi lain di jurusan, serta "
    "(4) peningkatan kemudahan pengelolaan sistem agar pembaruan data dapat berjalan "
    "secara lebih efisien."
)
doc.add_paragraph()

# Tanda tangan
doc.add_paragraph()
doc.add_paragraph()
p_ttd = doc.add_paragraph()
p_ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_ttd.add_run(f"Balikpapan, {datetime.date.today().strftime('%d %B %Y')}").font.size = Pt(11)

p_ttd2 = doc.add_paragraph()
p_ttd2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_ttd2.add_run("Kepala Laboratorium Inovasi Digital,").font.size = Pt(11)

for _ in range(4):
    p_sp = doc.add_paragraph()
    p_sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

p_nama = doc.add_paragraph()
p_nama.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_nama = p_nama.add_run("Aidil Saputra Kirsan")
run_nama.bold = True
run_nama.font.size = Pt(11)
run_nama.underline = True

# ─────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Dokumen berhasil dibuat: {OUTPUT_PATH}")
