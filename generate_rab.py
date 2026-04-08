"""
Generate RAB (Rencana Anggaran Biaya) DOCX untuk SITRIA - Galeri Inovasi Digital
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import locale
import os

def format_rupiah(value):
    """Format angka ke format Rupiah Indonesia"""
    return f"{value:,.0f}".replace(",", ".")

def set_cell_shading(cell, color):
    """Set background color of a table cell"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                          f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                          f'</w:tcBorders>')
    tcPr.append(tcBorders)

def set_row_height(row, height_cm):
    """Set row height"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(height_cm * 567)}" w:hRule="atLeast"/>')
    trPr.append(trHeight)

def add_cell_text(cell, text, bold=False, size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    """Helper to add formatted text to a cell"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def create_rab_document():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(33)  # F4/Folio
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # ========== KOP SURAT ==========
    kop = doc.add_paragraph()
    kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kop.add_run("KEMENTERIAN PENDIDIKAN, KEBUDAYAAN,\nRISET, DAN TEKNOLOGI")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    run.bold = True

    inst = doc.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = inst.add_run("INSTITUT TEKNOLOGI KALIMANTAN")
    run.font.size = Pt(13)
    run.font.name = "Times New Roman"
    run.bold = True

    alamat = doc.add_paragraph()
    alamat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = alamat.add_run(
        "Kampus ITK Karang Joang, Balikpapan 76127\n"
        "Telepon (0542) 8530801 Faksimile (0542) 8530800\n"
        "Surat elektronik: humas@itk.ac.id laman : www.itk.ac.id"
    )
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"

    # Garis pemisah
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = line._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # ========== JUDUL ==========
    doc.add_paragraph()  # spacing

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RENCANA ANGGARAN BIAYA (RAB)")
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.bold = True
    run.underline = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Pembuatan Aplikasi Web SITRIA \u2013 Galeri Inovasi Digital")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    run.bold = True

    unit = doc.add_paragraph()
    unit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = unit.add_run("Laboratorium Inovasi Digital \u2013 Institut Teknologi Kalimantan")
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"

    doc.add_paragraph()  # spacing

    # ========== INFO SECTION ==========
    info_items = [
        ("Jenis Belanja", ": Belanja Jasa Lainnya"),
        ("Metode Pengadaan", ": Penunjukan Langsung"),
        ("Dasar Hukum", ": Perpres No. 12 Tahun 2021 tentang Perubahan Atas Perpres No. 16\n                                    Tahun 2018 tentang Pengadaan Barang/Jasa Pemerintah"),
        ("Tahun Anggaran", ": 2026"),
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{label:<25}{value}")
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

    doc.add_paragraph()

    # ========== TABEL RAB ==========
    items = [
        # (no, nama, jumlah, satuan, harga_satuan, is_header, sub_items)
        # Section A - MODUL FRONTEND
        ("A", "MODUL FRONTEND", None, None, None, True, None),
        ("1", "Dashboard SINTA", 1, "paket", 3500000, False,
         "Statistik publikasi Scopus (Q1-Q4) & SINTA (S1-S6), H-Index, sitasi, "
         "visualisasi interaktif Chart.js & ECharts, mode TV Display auto-rotate untuk presentasi"),
        ("2", "Koleksi Karya / Galeri Inovasi", 1, "paket", 2500000, False,
         "6 kategori karya: Penelitian, Pengabdian, Scopus, SINTA, Buku, HKI; "
         "filter prodi, pencarian real-time, trend chart per tahun"),
        ("3", "Kolaborasi Riset", 1, "paket", 2000000, False,
         "Visualisasi klaster riset antar-prodi, deteksi potensi kolaborasi lintas program studi, "
         "badge level kolaborasi (High/Medium/Low)"),
        ("4", "Roadmap Riset", 1, "paket", 2000000, False,
         "Sankey Diagram evolusi topik riset 2018-2025, analisis temporal, "
         "deteksi tren riset berkelanjutan dan emerging"),
        ("5", "Expertise Finder / Cari Pakar", 1, "paket", 2000000, False,
         "Sistem pencarian pakar berdasarkan keahlian riset, ranking relevansi otomatis, "
         "instant search browser-based"),
        ("6", "Dashboard Dana & Hibah", 1, "paket", 2000000, False,
         "Analitik pendanaan riset & pengabdian, breakdown sumber dana internal/eksternal/BIMA, "
         "tren pendanaan per tahun, dukungan data LKPS & AMI"),
        ("7", "DTPS Akreditasi", 1, "paket", 1500000, False,
         "Data dosen tetap program studi untuk keperluan akreditasi dan pelaporan LKPS"),
        ("8", "UI/UX Design Premium", 1, "paket", 2000000, False,
         "Desain antarmuka glassmorphism, dark & light theme, responsive design, "
         "animasi transisi, layout multi-prodi"),

        # Section B - BACKEND & DATA PIPELINE
        ("B", "MODUL BACKEND & DATA PIPELINE", None, None, None, True, None),
        ("9", "Backend API Server", 1, "paket", 1500000, False,
         "Express.js server, proxy Google Scholar API, sistem caching, CORS management, "
         "environment configuration"),
        ("10", "Script Web Scraping SINTA", 1, "paket", 2500000, False,
         "Otomasi pengambilan data profil dosen, publikasi, penelitian, pengabdian, "
         "buku, dan HKI dari portal SINTA menggunakan Python + BeautifulSoup4"),
        ("11", "Script Algoritma AI \u2013 TF-IDF & K-Means Clustering", 1, "paket", 3000000, False,
         "Pengelompokan topik riset otomatis menggunakan scikit-learn, "
         "optimasi jumlah klaster via Silhouette Score, deteksi kolaborasi lintas-prodi"),
        ("12", "Script Algoritma AI \u2013 Expertise Indexing", 1, "paket", 2000000, False,
         "Pembangunan indeks keahlian dosen berbasis TF-IDF dari riwayat publikasi, "
         "ekstraksi keyword top-5 per profil dosen"),
        ("13", "Script Research Roadmap Analysis", 1, "paket", 1500000, False,
         "Analisis evolusi topik riset temporal untuk data Sankey diagram, "
         "deteksi topik berkelanjutan vs emerging topics"),

        # Section C - MANAGEMENT DATA
        ("C", "MODUL MANAGEMENT DATA", None, None, None, True, None),
        ("14", "Management Data Multi-Prodi", 1, "paket", 1500000, False,
         "Struktur data scalable untuk 10+ program studi, registry prodi, "
         "panduan penambahan data prodi baru"),
        ("15", "Management Data Dosen & Publikasi", 1, "paket", 1500000, False,
         "Pengelolaan data dosen, statistik SINTA, dokumen Scopus/SINTA per prodi, "
         "data penelitian, pengabdian, buku, dan HKI"),

        # Section D - DEPLOYMENT & DOKUMENTASI
        ("D", "DEPLOYMENT & DOKUMENTASI", None, None, None, True, None),
        ("16", "Deployment & Konfigurasi Server", 1, "paket", 1000000, False,
         "Setup production build Vite, konfigurasi environment, optimasi performa, "
         "deployment ke server kampus"),
        ("17", "Dokumentasi Teknis & Panduan Penggunaan", 1, "paket", 1000000, False,
         "Dokumentasi arsitektur sistem, alur data (data flow), panduan administrator, "
         "panduan penambahan prodi, dokumentasi API"),
    ]

    # Create table
    col_widths = [Cm(1), Cm(6.5), Cm(1.2), Cm(1.5), Cm(2.8), Cm(2.8)]

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["No.", "Nama Kegiatan", "Jumlah", "Satuan", "Harga\nSatuan (Rp)", "Harga\nTotal (Rp)"]
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        add_cell_text(cell, header, bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "D9E2F3")
        set_cell_border(cell)

    sub_total = 0

    for item in items:
        no, nama, jumlah, satuan, harga, is_header, sub_items = item
        row = table.add_row()

        if is_header:
            # Section header row
            cell_no = row.cells[0]
            add_cell_text(cell_no, no, bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(cell_no, "E8F0FE")
            set_cell_border(cell_no)

            cell_nama = row.cells[1]
            add_cell_text(cell_nama, nama, bold=True, size=9)
            set_cell_shading(cell_nama, "E8F0FE")
            set_cell_border(cell_nama)

            for i in range(2, 6):
                set_cell_shading(row.cells[i], "E8F0FE")
                set_cell_border(row.cells[i])
                row.cells[i].text = ""
        else:
            # Data row
            cell_no = row.cells[0]
            add_cell_text(cell_no, no, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_border(cell_no)

            # Nama kegiatan with sub-description
            cell_nama = row.cells[1]
            cell_nama.text = ""
            p = cell_nama.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(nama)
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            run.bold = True

            if sub_items:
                p2 = cell_nama.add_paragraph()
                p2.paragraph_format.space_before = Pt(1)
                p2.paragraph_format.space_after = Pt(2)
                run2 = p2.add_run(sub_items)
                run2.font.size = Pt(8)
                run2.font.name = "Times New Roman"
                run2.font.color.rgb = RGBColor(80, 80, 80)
            set_cell_border(cell_nama)

            # Jumlah
            cell_jml = row.cells[2]
            add_cell_text(cell_jml, str(jumlah), size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_border(cell_jml)

            # Satuan
            cell_sat = row.cells[3]
            add_cell_text(cell_sat, satuan, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_border(cell_sat)

            # Harga satuan
            cell_hs = row.cells[4]
            add_cell_text(cell_hs, format_rupiah(harga), size=9, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            set_cell_border(cell_hs)

            # Harga total
            total = jumlah * harga
            sub_total += total
            cell_ht = row.cells[5]
            add_cell_text(cell_ht, format_rupiah(total), size=9, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            set_cell_border(cell_ht)

    # Sub Total row
    row_sub = table.add_row()
    merged_cell = row_sub.cells[0].merge(row_sub.cells[4])
    add_cell_text(merged_cell, "Sub Total", bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(merged_cell, "F2F2F2")
    set_cell_border(merged_cell)
    cell_st = row_sub.cells[5]
    add_cell_text(cell_st, format_rupiah(sub_total), bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(cell_st, "F2F2F2")
    set_cell_border(cell_st)

    # PPN 11%
    ppn = sub_total * 0.11
    row_ppn = table.add_row()
    merged_ppn = row_ppn.cells[0].merge(row_ppn.cells[4])
    add_cell_text(merged_ppn, "PPN 11%", bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(merged_ppn, "F2F2F2")
    set_cell_border(merged_ppn)
    cell_ppn = row_ppn.cells[5]
    add_cell_text(cell_ppn, format_rupiah(ppn), bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(cell_ppn, "F2F2F2")
    set_cell_border(cell_ppn)

    # Total
    total_all = sub_total + ppn
    row_total = table.add_row()
    merged_total = row_total.cells[0].merge(row_total.cells[4])
    add_cell_text(merged_total, "TOTAL", bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(merged_total, "D9E2F3")
    set_cell_border(merged_total)
    cell_total = row_total.cells[5]
    add_cell_text(cell_total, format_rupiah(total_all), bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_shading(cell_total, "D9E2F3")
    set_cell_border(cell_total)

    # Set column widths
    for row in table.rows:
        for i, width in enumerate(col_widths):
            row.cells[i].width = width

    doc.add_paragraph()

    # ========== TERBILANG ==========
    p_terb = doc.add_paragraph()
    run = p_terb.add_run(f"Terbilang: ")
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run.bold = True
    run = p_terb.add_run("Tiga Puluh Tujuh Juta Seratus Delapan Puluh Lima Ribu Rupiah")
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run.italic = True

    doc.add_paragraph()

    # ========== CATATAN ==========
    p_cat_title = doc.add_paragraph()
    run = p_cat_title.add_run("Catatan:")
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run.bold = True
    run.underline = True

    catatan = [
        "Waktu pengerjaan maksimal 45 (Empat Puluh Lima) hari kalender terhitung mulai dari tanggal yang tercantum pada Surat Pesanan.",
        "Pembayaran dilakukan setelah proses serah terima produk dan produk diterima oleh Institut Teknologi Kalimantan.",
        "Aplikasi mencakup teknologi: Vue.js 3, Vite, Tailwind CSS, Chart.js, Apache ECharts, Express.js, Python (scikit-learn, BeautifulSoup4).",
        "Termasuk hak akses source code dan dokumentasi teknis lengkap.",
        "Garansi maintenance selama 30 (Tiga Puluh) hari kalender setelah serah terima.",
    ]

    for i, cat in enumerate(catatan, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{i}. {cat}")
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"

    doc.add_paragraph()
    doc.add_paragraph()

    # ========== TANDA TANGAN ==========
    # Two columns for signature
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left - Penyedia
    cell_left = sig_table.rows[0].cells[0]
    cell_left.text = ""
    p = cell_left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Penyedia,\n\n\n\n\n")
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run2 = p.add_run("(........................................)")
    run2.font.size = Pt(10)
    run2.font.name = "Times New Roman"

    # Right - PPK
    cell_right = sig_table.rows[0].cells[1]
    cell_right.text = ""
    p = cell_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Balikpapan, .................... 2026\nPejabat Pengadaan\nInstitut Teknologi Kalimantan\n\n\n\n")
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run2 = p.add_run("(........................................)")
    run2.font.size = Pt(10)
    run2.font.name = "Times New Roman"
    p3 = cell_right.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("NIP. ................................")
    run3.font.size = Pt(9)
    run3.font.name = "Times New Roman"

    # Save
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, "RAB_SITRIA_Galeri_Inovasi_Digital.docx")
    doc.save(output_path)
    print(f"Dokumen berhasil disimpan: {output_path}")
    print(f"\nRingkasan:")
    print(f"  Sub Total  : Rp {format_rupiah(sub_total)}")
    print(f"  PPN 11%    : Rp {format_rupiah(ppn)}")
    print(f"  TOTAL      : Rp {format_rupiah(total_all)}")

if __name__ == "__main__":
    create_rab_document()
